"""Integration test: document_reader → DocxParser full chain

Validates the complete read_document_as_text flow with real DOCX files.
No mocking of DocxParser — only the executor's file I/O is replaced
with real filesystem reads to simulate sandbox file access.
"""

from __future__ import annotations

import os
import tempfile

import pytest
from docx import Document
from docx.shared import Inches

from myrm_agent_harness.agent.meta_tools.file_ops.utils.document_reader import (
    is_document_path,
    read_document_as_text,
)


class _LocalFileExecutor:
    """Minimal executor stub that reads from real local filesystem."""

    async def read_file_bytes(self, path: str) -> bytes:
        if not os.path.exists(path):
            raise FileNotFoundError(f"No such file: {path}")
        with open(path, "rb") as f:
            return f.read()


@pytest.fixture
def executor() -> _LocalFileExecutor:
    return _LocalFileExecutor()


def _create_complex_docx(path: str) -> None:
    """Create a DOCX with headings, lists, tables (including merged cells)."""
    doc = Document()
    doc.add_heading("Project Report", level=1)
    doc.add_paragraph("This report summarizes Q2 progress.")

    doc.add_heading("Key Items", level=2)
    doc.add_paragraph("Complete onboarding", style="List Bullet")
    doc.add_paragraph("Ship v2.0", style="List Bullet")

    doc.add_heading("Milestones", level=2)
    doc.add_paragraph("First milestone", style="List Number")
    doc.add_paragraph("Second milestone", style="List Number")

    doc.add_heading("Budget Table", level=2)
    table = doc.add_table(rows=3, cols=3)
    table.cell(0, 0).text = "Category"
    table.cell(0, 1).text = "Q1"
    table.cell(0, 2).text = "Q2"
    table.cell(1, 0).text = "Engineering"
    table.cell(1, 1).text = "100k"
    table.cell(1, 2).text = "120k"
    table.cell(2, 0).text = "Marketing"
    table.cell(2, 1).text = "50k"
    table.cell(2, 2).text = "60k"

    doc.add_paragraph("End of report.")

    # Table with merged cells
    doc.add_heading("Merged Table", level=2)
    t2 = doc.add_table(rows=2, cols=3)
    t2.cell(0, 0).text = "Span"
    t2.cell(0, 1).text = "B"
    t2.cell(0, 2).text = "C"
    t2.cell(1, 0).text = "X"
    t2.cell(1, 1).text = "Y"
    t2.cell(1, 2).text = "Z"
    t2.cell(0, 0).merge(t2.cell(0, 1))

    doc.save(path)


class TestDocumentReaderDocxIntegration:
    """Full integration: read_document_as_text with real DOCX parsing."""

    @pytest.mark.asyncio
    async def test_full_chain_headings_and_paragraphs(self, executor: _LocalFileExecutor) -> None:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            _create_complex_docx(f.name)
            tmp = f.name

        try:
            result = await read_document_as_text(tmp, executor)
            assert "[Document:" in result
            assert "# Project Report" in result
            assert "## Key Items" in result
            assert "This report summarizes Q2 progress." in result
        finally:
            os.unlink(tmp)

    @pytest.mark.asyncio
    async def test_full_chain_lists(self, executor: _LocalFileExecutor) -> None:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            _create_complex_docx(f.name)
            tmp = f.name

        try:
            result = await read_document_as_text(tmp, executor)
            assert "- Complete onboarding" in result
            assert "- Ship v2.0" in result
            assert "1. First milestone" in result
            assert "1. Second milestone" in result
        finally:
            os.unlink(tmp)

    @pytest.mark.asyncio
    async def test_full_chain_table(self, executor: _LocalFileExecutor) -> None:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            _create_complex_docx(f.name)
            tmp = f.name

        try:
            result = await read_document_as_text(tmp, executor)
            assert "| Category | Q1 | Q2 |" in result
            assert "| Engineering | 100k | 120k |" in result
            assert "| Marketing | 50k | 60k |" in result
        finally:
            os.unlink(tmp)

    @pytest.mark.asyncio
    async def test_full_chain_merged_cells(self, executor: _LocalFileExecutor) -> None:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            _create_complex_docx(f.name)
            tmp = f.name

        try:
            result = await read_document_as_text(tmp, executor)
            # In the merged table, "Span" appears only once, merged cell is empty
            merged_section = result[result.index("## Merged Table"):]
            table_lines = [l for l in merged_section.split("\n") if "|" in l and "---" not in l]
            header_cells = [c.strip() for c in table_lines[0].split("|")[1:-1]]
            assert header_cells[0] != ""
            assert header_cells[1] == ""
            assert header_cells[2] == "C"
        finally:
            os.unlink(tmp)

    @pytest.mark.asyncio
    async def test_full_chain_document_order(self, executor: _LocalFileExecutor) -> None:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            _create_complex_docx(f.name)
            tmp = f.name

        try:
            result = await read_document_as_text(tmp, executor)
            # Verify overall order: heading → paragraph → lists → table → end text
            idx_heading = result.index("# Project Report")
            idx_summary = result.index("This report summarizes")
            idx_bullet = result.index("- Complete onboarding")
            idx_table = result.index("| Category")
            idx_end = result.index("End of report")
            assert idx_heading < idx_summary < idx_bullet < idx_table < idx_end
        finally:
            os.unlink(tmp)

    @pytest.mark.asyncio
    async def test_file_not_found_propagates(self, executor: _LocalFileExecutor) -> None:
        with pytest.raises(FileNotFoundError):
            await read_document_as_text("/nonexistent/file.docx", executor)

    @pytest.mark.asyncio
    async def test_empty_document_message(self, executor: _LocalFileExecutor) -> None:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            doc.save(f.name)
            tmp = f.name

        try:
            result = await read_document_as_text(tmp, executor)
            assert "No extractable content" in result
        finally:
            os.unlink(tmp)


class TestDocumentReaderTruncation:
    """Test large document truncation behavior."""

    @pytest.mark.asyncio
    async def test_large_document_truncated(self, executor: _LocalFileExecutor) -> None:
        """Documents exceeding 200k chars should be truncated with marker."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            # Each paragraph ~100 chars, need >2000 paragraphs for 200k
            for i in range(2500):
                doc.add_paragraph(f"Paragraph number {i} " + "x" * 80)
            doc.save(f.name)
            tmp = f.name

        try:
            result = await read_document_as_text(tmp, executor)
            assert "truncated at 200000 chars" in result
        finally:
            os.unlink(tmp)


class TestDocumentReaderEdgeCases:
    """Edge cases for the document_reader integration layer."""

    @pytest.mark.asyncio
    async def test_unsupported_suffix_returns_message(self, executor: _LocalFileExecutor) -> None:
        """Unsupported document format should return descriptive message."""
        with tempfile.NamedTemporaryFile(suffix=".odt", delete=False) as f:
            f.write(b"fake content")
            tmp = f.name

        try:
            result = await read_document_as_text(tmp, executor)
            assert "Unsupported document format" in result
        finally:
            os.unlink(tmp)

    @pytest.mark.asyncio
    async def test_corrupted_docx_returns_error(self, executor: _LocalFileExecutor) -> None:
        """Corrupted DOCX file should return a parsing error message, not crash."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            f.write(b"This is not a valid DOCX file at all")
            tmp = f.name

        try:
            result = await read_document_as_text(tmp, executor)
            assert "Parsing failed" in result
        finally:
            os.unlink(tmp)


class TestIsDocumentPath:
    """Integration: is_document_path utility function."""

    def test_docx_detected(self) -> None:
        assert is_document_path("/workspace/report.docx") is True

    def test_non_document_not_detected(self) -> None:
        assert is_document_path("/workspace/script.py") is False

    def test_xlsx_detected(self) -> None:
        assert is_document_path("data.xlsx") is True

    def test_pptx_detected(self) -> None:
        assert is_document_path("/slides/deck.pptx") is True

    def test_ipynb_detected(self) -> None:
        assert is_document_path("notebook.ipynb") is True

    def test_case_insensitive_extension(self) -> None:
        assert is_document_path("Report.DOCX") is True  # .suffix.lower() handles case
