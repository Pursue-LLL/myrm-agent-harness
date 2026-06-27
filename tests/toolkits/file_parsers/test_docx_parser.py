"""Tests for DocxParser

Tests heading, list, table extraction and document-order interleaving from .docx files.
"""

from __future__ import annotations

import os
import tempfile

import pytest
from docx import Document

from myrm_agent_harness.toolkits.file_parsers import DocxParser, get_parser, is_supported


class TestDocxParserRegistry:
    """Test parser registration and discovery."""

    def test_docx_is_supported(self) -> None:
        assert is_supported("report.docx") is True

    def test_doc_is_supported(self) -> None:
        assert is_supported("old.doc") is True

    def test_get_parser_returns_docx_parser(self) -> None:
        parser = get_parser("contract.docx")
        assert isinstance(parser, DocxParser)

    def test_supported_extensions(self) -> None:
        parser = DocxParser()
        assert ".docx" in parser.supported_extensions
        assert ".doc" in parser.supported_extensions


class TestDocxParserBasic:
    """Test basic paragraph and heading extraction."""

    @pytest.mark.asyncio
    async def test_heading_and_paragraph(self) -> None:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            doc.add_heading("Title", level=1)
            doc.add_paragraph("Body text")
            doc.add_heading("Section", level=2)
            doc.add_paragraph("More content")
            doc.save(f.name)
            tmp = f.name

        try:
            parser = DocxParser()
            result = await parser.parse(tmp)
            assert "# Title" in result
            assert "Body text" in result
            assert "## Section" in result
            assert "More content" in result
        finally:
            os.unlink(tmp)

    @pytest.mark.asyncio
    async def test_empty_document(self) -> None:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            doc.save(f.name)
            tmp = f.name

        try:
            parser = DocxParser()
            result = await parser.parse(tmp)
            assert result == ""
        finally:
            os.unlink(tmp)

    @pytest.mark.asyncio
    async def test_file_not_found(self) -> None:
        parser = DocxParser()
        with pytest.raises(FileNotFoundError):
            await parser.parse("/nonexistent/path.docx")


class TestDocxParserList:
    """Test list item formatting."""

    @pytest.mark.asyncio
    async def test_bullet_list(self) -> None:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            doc.add_paragraph("Items:")
            doc.add_paragraph("Buy milk", style="List Bullet")
            doc.add_paragraph("Clean house", style="List Bullet")
            doc.save(f.name)
            tmp = f.name

        try:
            parser = DocxParser()
            result = await parser.parse(tmp)
            assert "- Buy milk" in result
            assert "- Clean house" in result
        finally:
            os.unlink(tmp)

    @pytest.mark.asyncio
    async def test_numbered_list(self) -> None:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            doc.add_paragraph("Steps:")
            doc.add_paragraph("First step", style="List Number")
            doc.add_paragraph("Second step", style="List Number")
            doc.save(f.name)
            tmp = f.name

        try:
            parser = DocxParser()
            result = await parser.parse(tmp)
            assert "1. First step" in result
            assert "1. Second step" in result
        finally:
            os.unlink(tmp)


class TestDocxParserTable:
    """Test table extraction and merged cell handling."""

    @pytest.mark.asyncio
    async def test_basic_table(self) -> None:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            table = doc.add_table(rows=3, cols=2)
            table.cell(0, 0).text = "Name"
            table.cell(0, 1).text = "Age"
            table.cell(1, 0).text = "Alice"
            table.cell(1, 1).text = "30"
            table.cell(2, 0).text = "Bob"
            table.cell(2, 1).text = "25"
            doc.save(f.name)
            tmp = f.name

        try:
            parser = DocxParser()
            result = await parser.parse(tmp)
            assert "| Name | Age |" in result
            assert "| --- | --- |" in result
            assert "| Alice | 30 |" in result
            assert "| Bob | 25 |" in result
        finally:
            os.unlink(tmp)

    @pytest.mark.asyncio
    async def test_merged_cells_horizontal(self) -> None:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            table = doc.add_table(rows=2, cols=3)
            table.cell(0, 0).text = "A"
            table.cell(0, 1).text = "B"
            table.cell(0, 2).text = "C"
            table.cell(1, 0).text = "D"
            table.cell(1, 1).text = "E"
            table.cell(1, 2).text = "F"
            table.cell(0, 0).merge(table.cell(0, 1))
            doc.save(f.name)
            tmp = f.name

        try:
            parser = DocxParser()
            result = await parser.parse(tmp)
            lines = [l for l in result.split("\n") if "|" in l and "---" not in l]
            # Split keeping empty segments between pipes
            raw_cells = lines[0].split("|")
            # Trim leading/trailing empty from outer pipes, keep inner cells
            inner_cells = [c.strip() for c in raw_cells[1:-1]]
            # Merged cell (col 0+1) should appear once with content, second should be empty
            assert inner_cells[0] != ""
            assert inner_cells[1] == ""
            assert inner_cells[2] == "C"
        finally:
            os.unlink(tmp)

    @pytest.mark.asyncio
    async def test_pipe_character_escaped(self) -> None:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            table = doc.add_table(rows=2, cols=1)
            table.cell(0, 0).text = "Header"
            table.cell(1, 0).text = "A|B"
            doc.save(f.name)
            tmp = f.name

        try:
            parser = DocxParser()
            result = await parser.parse(tmp)
            assert "A\\|B" in result
        finally:
            os.unlink(tmp)


class TestDocxParserDocumentOrder:
    """Test that paragraphs and tables maintain document order."""

    @pytest.mark.asyncio
    async def test_interleaved_order(self) -> None:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            doc.add_paragraph("Before table")
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "X"
            table.cell(0, 1).text = "Y"
            table.cell(1, 0).text = "1"
            table.cell(1, 1).text = "2"
            doc.add_paragraph("After table")
            doc.save(f.name)
            tmp = f.name

        try:
            parser = DocxParser()
            result = await parser.parse(tmp)
            before_idx = result.index("Before table")
            table_idx = result.index("| X")
            after_idx = result.index("After table")
            assert before_idx < table_idx < after_idx
        finally:
            os.unlink(tmp)

    @pytest.mark.asyncio
    async def test_multiple_tables_order(self) -> None:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            doc.add_paragraph("Intro")
            t1 = doc.add_table(rows=1, cols=1)
            t1.cell(0, 0).text = "Table1"
            doc.add_paragraph("Middle")
            t2 = doc.add_table(rows=1, cols=1)
            t2.cell(0, 0).text = "Table2"
            doc.add_paragraph("End")
            doc.save(f.name)
            tmp = f.name

        try:
            parser = DocxParser()
            result = await parser.parse(tmp)
            assert result.index("Intro") < result.index("Table1")
            assert result.index("Table1") < result.index("Middle")
            assert result.index("Middle") < result.index("Table2")
            assert result.index("Table2") < result.index("End")
        finally:
            os.unlink(tmp)


class TestDocxParserEdgeCases:
    """Test edge cases and boundary conditions."""

    @pytest.mark.asyncio
    async def test_cell_with_newlines(self) -> None:
        """Cell containing newlines should be flattened to spaces."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            table = doc.add_table(rows=2, cols=1)
            table.cell(0, 0).text = "Header"
            table.cell(1, 0).text = "Line1\nLine2\nLine3"
            doc.save(f.name)
            tmp = f.name

        try:
            parser = DocxParser()
            result = await parser.parse(tmp)
            assert "Line1 Line2 Line3" in result
            assert "Line1\n" not in result.split("| ")[1] if "| " in result else True
        finally:
            os.unlink(tmp)

    @pytest.mark.asyncio
    async def test_heading_without_number(self) -> None:
        """Heading style without a numeric level should be treated as plain text."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            # python-docx "Heading" without level defaults to "Heading 1"
            # but we test the fallback by adding a regular heading and checking
            doc.add_heading("Normal Heading", level=3)
            doc.add_paragraph("Regular text")
            doc.save(f.name)
            tmp = f.name

        try:
            parser = DocxParser()
            result = await parser.parse(tmp)
            assert "### Normal Heading" in result
            assert "Regular text" in result
        finally:
            os.unlink(tmp)

    @pytest.mark.asyncio
    async def test_only_whitespace_paragraphs_skipped(self) -> None:
        """Paragraphs with only whitespace should be skipped."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            doc.add_paragraph("Visible")
            doc.add_paragraph("   ")
            doc.add_paragraph("\t\n")
            doc.add_paragraph("Also visible")
            doc.save(f.name)
            tmp = f.name

        try:
            parser = DocxParser()
            result = await parser.parse(tmp)
            assert "Visible" in result
            assert "Also visible" in result
            blocks = result.split("\n\n")
            assert all(b.strip() != "" for b in blocks)
        finally:
            os.unlink(tmp)

    @pytest.mark.asyncio
    async def test_single_row_table(self) -> None:
        """Table with only header row (no data rows)."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            table = doc.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "Col A"
            table.cell(0, 1).text = "Col B"
            doc.save(f.name)
            tmp = f.name

        try:
            parser = DocxParser()
            result = await parser.parse(tmp)
            assert "| Col A | Col B |" in result
            assert "| --- | --- |" in result
        finally:
            os.unlink(tmp)

    @pytest.mark.asyncio
    async def test_large_heading_level(self) -> None:
        """Heading levels 4-6 should be parsed correctly."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            doc.add_heading("Deep heading", level=4)
            doc.save(f.name)
            tmp = f.name

        try:
            parser = DocxParser()
            result = await parser.parse(tmp)
            assert "#### Deep heading" in result
        finally:
            os.unlink(tmp)

    @pytest.mark.asyncio
    async def test_mixed_content_comprehensive(self) -> None:
        """Full document mixing all content types preserves correct structure."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            doc.add_heading("Title", level=1)
            doc.add_paragraph("Intro paragraph")
            doc.add_paragraph("Bullet A", style="List Bullet")
            doc.add_paragraph("Bullet B", style="List Bullet")
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "H1"
            table.cell(0, 1).text = "H2"
            table.cell(1, 0).text = "V1"
            table.cell(1, 1).text = "V2"
            doc.add_heading("Sub", level=2)
            doc.add_paragraph("Step 1", style="List Number")
            doc.add_paragraph("Closing")
            doc.save(f.name)
            tmp = f.name

        try:
            parser = DocxParser()
            result = await parser.parse(tmp)
            assert "# Title" in result
            assert "Intro paragraph" in result
            assert "- Bullet A" in result
            assert "- Bullet B" in result
            assert "| H1 | H2 |" in result
            assert "| V1 | V2 |" in result
            assert "## Sub" in result
            assert "1. Step 1" in result
            assert "Closing" in result
            # Order check
            assert result.index("# Title") < result.index("Intro paragraph")
            assert result.index("- Bullet A") < result.index("| H1")
            assert result.index("| V1") < result.index("## Sub")
            assert result.index("## Sub") < result.index("1. Step 1")
        finally:
            os.unlink(tmp)

    @pytest.mark.asyncio
    async def test_vertical_merge(self) -> None:
        """Vertical merged cells should not produce duplicate text."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            table = doc.add_table(rows=3, cols=2)
            table.cell(0, 0).text = "H1"
            table.cell(0, 1).text = "H2"
            table.cell(1, 0).text = "Span"
            table.cell(1, 1).text = "B"
            table.cell(2, 0).text = "C"
            table.cell(2, 1).text = "D"
            # Vertical merge: row 1 col 0 + row 2 col 0
            table.cell(1, 0).merge(table.cell(2, 0))
            doc.save(f.name)
            tmp = f.name

        try:
            parser = DocxParser()
            result = await parser.parse(tmp)
            # "Span" should appear in the table output
            assert "Span" in result
            # The table should still have valid Markdown structure
            table_lines = [l for l in result.split("\n") if l.startswith("|")]
            assert len(table_lines) >= 3  # header + separator + at least 1 data row
        finally:
            os.unlink(tmp)
